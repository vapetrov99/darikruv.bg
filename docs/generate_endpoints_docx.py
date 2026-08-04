"""Generate endpoints table DOCX for diploma section 8."""
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT = Path(__file__).resolve().parent / "8-endpoints-pregled.docx"

ROWS = [
    ("1", "POST", "register", "Публичен", "users, donors", "Регистрация; hash на парола; имейл за верификация; при дарител — запис в donors"),
    ("2", "GET", "verify_email", "Публичен", "users", "Потвърждение на имейл по token; is_verified=1"),
    ("3", "POST", "login", "Публичен", "users, donors, rate_limit_attempts", "Вход; проверка на парола; JWT; изисква потвърден имейл"),
    ("4", "POST", "request_password_reset", "Публичен", "users, rate_limit_attempts", "Заявка за забравена парола; имейл с линк (60 мин)"),
    ("5", "POST", "reset_password", "Публичен", "users", "Нова парола с валиден reset token"),
    ("6", "GET", "requests", "Опционален JWT", "blood_requests, request_responses", "Списък активни заявки (48h); при login — статус на отзов"),
    ("7", "GET", "request_details", "Опционален JWT", "blood_requests, request_responses", "Детайли за една заявка по id"),
    ("8", "POST", "create_request", "JWT", "blood_requests", "Нова заявка; push + email опашка към дарители"),
    ("9", "POST", "update_request", "JWT (автор)", "blood_requests", "Редакция на собствена активна заявка (до 72h)"),
    ("10", "POST", "respond_to_request", "JWT", "request_responses, blood_requests", "Отзов: pledge (24h) / confirm (потвърждение)"),
    ("11", "GET", "request_comments", "Публичен", "request_comments", "Коментари към заявка по request_id"),
    ("12", "POST", "create_request_comment", "Публичен*", "request_comments", "Нов коментар (name, text); в UI — след login"),
    ("13", "GET", "my_requests", "JWT", "blood_requests", "Заявки, създадени от текущия user"),
    ("14", "GET", "my_responses", "JWT", "request_responses, blood_requests", "История на отзови на дарителя"),
    ("15", "POST", "update_profile", "JWT", "users, donors", "Редакция на профил; при donor — кръвна група, известия"),
    ("16", "POST", "update_last_donation", "JWT", "donors", "Дата на последно даряване (календар)"),
    ("17", "POST", "delete_account", "JWT", "много таблици", "Изтриване на акаунт и свързани данни"),
    ("18", "GET", "push_public_config", "Публичен", "—", "Публични Firebase/VAPID настройки за push"),
    ("19", "POST", "save_push_token", "JWT (donor)", "donor_push_tokens", "Запис на FCM token за push известия"),
    ("20", "GET", "users", "Admin", "users", "Списък потребители (админ панел)"),
    ("21", "GET", "donors", "Admin", "donors, users", "Списък дарители (админ панел)"),
    ("22", "POST", "create_campaign", "Admin", "donors", "Имейл известия за нова кампания"),
    ("23", "POST", "process_email_queue", "Admin", "email_queue", "Обработка на опашката за имейл известия"),
    ("24", "GET", "ncth_stores", "Публичен", "— (proxy)", "Пунктове за даряване от НЦТХ (външен API)"),
]

HEADERS = ["№", "Метод", "route", "Автентикация", "Таблици", "Функционалност"]


def main() -> None:
    doc = Document()

    title = doc.add_heading("8. Преглед на създадените крайни точки (endpoints)", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph(
        "Всички заявки към backend-а минават през единен вход: "
        "api/index.php?route=<име>, комбиниран с HTTP метод. Отговорът е в JSON формат."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=len(HEADERS))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, text in enumerate(HEADERS):
        p = hdr_cells[i].paragraphs[0]
        p.text = text
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)

    for row_data in ROWS:
        row = table.add_row().cells
        for i, text in enumerate(row_data):
            p = row[i].paragraphs[0]
            p.text = text
            for run in p.runs:
                run.font.size = Pt(8)

    doc.add_paragraph()
    legend = doc.add_heading("Легенда — автентикация", level=2)
    legend.alignment = WD_ALIGN_PARAGRAPH.LEFT

    legends = [
        "Публичен — без Authorization header",
        "JWT — Authorization: Bearer <token> (auth_require_user)",
        "Admin — JWT + роля admin",
        "Опционален JWT — работи без token; с token добавя допълнителни данни",
    ]
    for item in legends:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    note = doc.add_paragraph(
        "Бележка: create_request_comment е публичен в API; в интерфейса се ползва от страница "
        "request-details.html, която изисква вход (requireAuth)."
    )
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(9)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    doc.save(str(OUTPUT))
    print("Created:", str(OUTPUT))


if __name__ == "__main__":
    main()
